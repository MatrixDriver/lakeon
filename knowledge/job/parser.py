"""Document parsing: PDF (pymupdf4llm), DOCX (python-docx), Markdown (direct read)."""
import logging
from typing import List, Dict, Tuple

logger = logging.getLogger(__name__)

def parse_document(file_path: str, format: str) -> Tuple[str, List[Dict]]:
    """Parse a document and return (markdown_text, page_metadata).

    page_metadata: [{page: int, char_start: int, char_end: int}, ...]
    For non-PDF formats, page_metadata is empty.
    """
    format = format.upper()
    if format == "PDF":
        return _parse_pdf(file_path)
    elif format == "DOCX":
        return _parse_docx(file_path), []
    elif format == "EPUB":
        return _parse_epub(file_path), []
    elif format in ("MARKDOWN", "TEXT"):
        return _parse_markdown(file_path), []
    else:
        raise ValueError(f"Unsupported format: {format}")

def _parse_pdf(file_path: str) -> Tuple[str, List[Dict]]:
    import pymupdf4llm
    import pymupdf

    # Get per-page markdown to build page_metadata
    doc = pymupdf.open(file_path)
    page_count = len(doc)
    doc.close()

    # Use pymupdf4llm with page_chunks to get per-page content
    pages = pymupdf4llm.to_markdown(file_path, page_chunks=True)

    page_metadata = []
    parts = []
    current_offset = 0

    for i, page_data in enumerate(pages):
        page_num = page_data.get("metadata", {}).get("page", 0)
        text = page_data.get("text", "")
        if not text:
            continue
        char_start = current_offset
        parts.append(text)
        current_offset += len(text)
        # Add separator between pages
        if i < len(pages) - 1:
            parts.append("\n\n")
            current_offset += 2
        page_metadata.append({
            "page": page_num + 1,  # 1-based page numbers
            "char_start": char_start,
            "char_end": current_offset,
        })

    markdown = "".join(parts)
    return markdown, page_metadata

def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "Heading 1" in style:
            parts.append(f"# {text}")
        elif "Heading 2" in style:
            parts.append(f"## {text}")
        elif "Heading 3" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            parts.append(rows[0])
            parts.append(header_sep)
            parts.extend(rows[1:])
    return "\n\n".join(parts)

def _parse_epub(file_path: str) -> str:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(file_path, options={"ignore_ncx": True})
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        # Convert headings to markdown
        for tag in soup.find_all(["h1", "h2", "h3", "h4"]):
            level = int(tag.name[1])
            tag.replace_with(f"\n{'#' * level} {tag.get_text().strip()}\n")
        text = soup.get_text(separator="\n").strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _parse_markdown(file_path: str) -> str:
    for encoding in ["utf-8", "gbk", "gb2312", "latin1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # Last resort: read as bytes and decode with replacement
    with open(file_path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")
